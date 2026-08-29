import io
import re
import os
import csv
from typing import Dict, Any, Tuple
import pypdf
import docx
import openpyxl

def clean_extracted_text(text: str) -> str:
    """Normalizes extracted text, removes extraneous control characters and redundant whitespace."""
    if not text:
        return ""
    # Normalize line breaks
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse multiple consecutive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse excessive horizontal spaces
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()

def parse_pdf_bytes(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    """Extracts text from PDF bytes using pypdf."""
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    num_pages = len(reader.pages)
    page_texts = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text.strip():
            page_texts.append(f"--- Page {i + 1} ---\n{page_text}")
    return "\n\n".join(page_texts), {"num_pages": num_pages}

def parse_docx_bytes(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    """Extracts text and tables from DOCX bytes using python-docx."""
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # Also extract table cells
    table_texts = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(row_data))
        if rows:
            table_texts.append("\n".join(rows))
            
    combined = "\n\n".join(paragraphs)
    if table_texts:
        combined += "\n\n--- Tables Extracted ---\n" + "\n\n".join(table_texts)
    return combined, {"paragraphs": len(paragraphs), "tables": len(doc.tables)}

def parse_xlsx_bytes(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    """Extracts tabular data from XLSX workbook sheets using openpyxl."""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    sheet_texts = []
    for sheetname in wb.sheetnames:
        sheet = wb[sheetname]
        rows = []
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                rows.append(row_str)
        if rows:
            sheet_texts.append(f"--- Sheet: {sheetname} ---\n" + "\n".join(rows))
    return "\n\n".join(sheet_texts), {"sheets": wb.sheetnames}

def parse_csv_bytes(file_bytes: bytes) -> Tuple[str, Dict[str, Any]]:
    """Extracts text from CSV bytes."""
    decoded = file_bytes.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(decoded))
    lines = [" | ".join(row) for row in reader if row]
    return "\n".join(lines), {"lines": len(lines)}

def parse_document_buffer(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Universal document parser supporting PDF, DOCX, XLSX, CSV, TXT, and Markdown.
    """
    ext = os.path.splitext(filename)[1].lower()
    extracted_text = ""
    metadata: Dict[str, Any] = {}

    try:
        if ext == ".pdf":
            extracted_text, metadata = parse_pdf_bytes(file_bytes)
        elif ext in [".docx", ".doc"]:
            extracted_text, metadata = parse_docx_bytes(file_bytes)
        elif ext in [".xlsx", ".xls"]:
            extracted_text, metadata = parse_xlsx_bytes(file_bytes)
        elif ext == ".csv":
            extracted_text, metadata = parse_csv_bytes(file_bytes)
        elif ext in [".txt", ".md", ".json", ".rtf"]:
            extracted_text = file_bytes.decode("utf-8", errors="ignore")
        else:
            extracted_text = file_bytes.decode("utf-8", errors="ignore")

        cleaned = clean_extracted_text(extracted_text)
        if not cleaned:
            raise ValueError("Could not extract readable text from document. File may be empty or encrypted.")

        return {
            "success": True,
            "filename": filename,
            "fileType": ext.replace(".", "").upper(),
            "extractedText": cleaned,
            "metadata": metadata
        }
    except Exception as e:
        raise RuntimeError(f"Failed to parse document: {str(e)}")
